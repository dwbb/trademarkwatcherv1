import yagmail
import docx
from termcolor import colored
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import os
from datetime import date, datetime, timedelta
import tess_scraper_functions as tess
from tess_scraper_functions import days_ago

### Will focus on Twitter API for now. MIght be simpler to work on format for emails and send emails manually through a lsitserv.
##ROUGH ATTEMPT! RETRY WITH GMAIL API instead!!!!!

def get_date():
# Makes sure date matched other py files, File date is current date and search date is the date used to search TESS

    file_date = tess.tess_search_date()

    search_date = date.today() - timedelta(days = days_ago)
    search_date = search_date.strftime("%m/%d/%Y")

    return search_date, file_date

def make_doc(content):

    search_date, file_date = get_date()

    # Create document
    doc = docx.Document()

    # Add a header to the document
    heading = doc.sections[0].header
    head = heading.paragraphs[0]
    head.text = "\tTrademark Database Surveillance Entries\t"

    entry_counter = 1

    if isinstance(content, list):
        if content == []:
            doc.add_paragraph('No results found for ' + str(search_date) + "...")
        else:
            for entry in content:
                #Number entries - number is bolded and centered
                num = "(" + str(entry_counter) + ")"
                p = doc.add_paragraph()

                paragraph_format = p.paragraph_format
                runner = p.add_run(num)
                runner.bold = True
                paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

                content = doc.add_paragraph("\n" + entry + '\n')


                entry_counter += 1
    else:
        doc.add_paragraph(content)

    filename = 'Tess Log (' + str(file_date) + ').docx'

    doc.save(filename)

    file_path = os.path.abspath(filename)

    return file_path


def send_email(content, sender = "***", recipient_list = "***", email_password = "***"):
# Sends email from sender to recipient_list emails

    attachment = make_doc(content)
    # Sender and recipient information
    sender_email = sender
    reciever_emails = recipient_list
    sender_password = email_password

    search_date, file_date = get_date()

    ## Email Subject
    subject = "Trademark Watcher " + search_date

    #Email body
    template_data = ["***"]
    email_template = "Hello {name},\n\nAs always see attached for this week's round-up of surveillance-related registries on the trademark database.\n\nWe are always seeking to improve this process and better meet the needs of anyone who finds this useful, so please do not hesitate to email us with any advice or feedback.\n\nStay vigilant,\nTrademark Watcher\nhttps://twitter.com/Trademark_Watch\n\n"

    //yag = yagmail.SMTP(user="***", password = "***")

    #Send Email
    for name in template_data:
        text = email_template.format(name=name)
        contents = [text, attachment]
        yag.send(reciever_emails, subject, contents)

    return print("Email sent.")
